from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from django.core.validators import URLValidator
from django.core.exceptions import ValidationError
import docx
from docx import *
import time

# data from user

text = ""
file_path = ""
document_name = ""
generate_definitions_table = False
panopto_link = ""
login = ""
password = ""


def get_basic_data_from_user():
    global file_path, document_name, generate_definitions_table, panopto_link, login, password
    validate = URLValidator()
    doc = docx.Document()

    link_is_valid = False
    while not link_is_valid:
        panopto_link = input("Please paste link to the Panopto video: ")
        try:
            validate(panopto_link)
            if panopto_link.__contains__("panopto.eu"):
                link_is_valid = True
            else:
                print("not a Panopto link.")
        except ValidationError:
            print("Incorrect link.")

    document_can_be_created = False
    while not document_can_be_created:
        document_name = input("Please enter the name of the word file you wish to create: ")
        file_path = input("...and directory path (where you want to create it): ")
        try:
            doc.add_heading(document_name, 0)
            doc.save(file_path + "\\" + document_name + ".docx")
            document_can_be_created = True
        except:
            print("incorrect path or file name.")

    login = input("Please enter your uni's email: ")
    password = input("Please enter your uni's account password: ")


def microsoft_login(driver):
    global login, password
    try:
        email_textbox = driver.find_element_by_name("loginfmt")
        email_textbox.send_keys(login)
        email_textbox.send_keys(Keys.RETURN)
        # goes to password page
        password_textbox = driver.find_element_by_name("passwd")
        password_textbox.send_keys(password)
        time.sleep(2)
        password_button = driver.find_element_by_id("idSIButton9")
        password_button.send_keys(Keys.RETURN)
        # do you want to remember next time: no
        not_remember_button = driver.find_element_by_id("idBtn_Back")
        not_remember_button.send_keys(Keys.RETURN)
    except:
        print("wrong login or password.")


def get_captions(driver):
    global text
    time.sleep(3)
    tab_header = driver.find_element_by_id("transcriptTabHeader")
    tab_header.click()
    time.sleep(2)
    list_of_all = driver.find_element_by_id("eventTabPanes").find_element_by_id(
        "transcriptTabPane").find_element_by_class_name("event-tab-scroll-pane").find_element_by_class_name(
        "event-tab-list")
    captions_divs = list_of_all.find_elements_by_class_name("index-event")
    captions_divs_rows = []
    for div_row in captions_divs:
        captions_divs_rows.append(div_row.find_element_by_class_name("index-event-row"))
    captions_text_spans = []
    for index_row in captions_divs_rows:
        captions_text_spans.append(index_row.find_element_by_class_name("event-text"))
    for span in captions_text_spans:
        text += (span.find_element_by_tag_name("span").text + "\n")
    print(text)


def add_text_to_word_file(heading, text):
    document = Document(file_path + "\\" + document_name + ".docx")
    document.add_heading(heading, 1)
    document.add_paragraph(text)
    document.save(file_path + "\\" + document_name + ".docx")



get_basic_data_from_user()


# selenium setup
PATH = "C:\Program Files (x86)\chromedriver.exe"
driver = webdriver.Chrome(PATH)  # PATH is location of chrome web driver needed to use selenium
driver.get(panopto_link)
login_button = driver.find_element_by_id("PageContentPlaceholder_loginControl_externalLoginButton")
login_button.send_keys(Keys.RETURN)  # press enter
# now, microsoft login page pops out
microsoft_login(driver)
# panopto video page
get_captions(driver)
# add transcript
add_text_to_word_file("Captions transcript", text)